import random
import os
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.font import Font


class GenerateurQCM:
    """Classe pour générer des QCM personnalisés à partir d'un fichier de questions."""
    
    def __init__(self):
        """Initialise le générateur de QCM."""
        self.questions = []
        self.fichier_entree = ""
        self.nombre_eleves = 0
        self.nombre_questions = 0
        self.graine = 0
        
    def lire_fichier_qcm(self, nom_fichier):
        """
        Lit le fichier contenant les questions et réponses.
        
        Args:
            nom_fichier (str): Chemin vers le fichier de questions.
            
        Returns:
            tuple: (succès (bool), message (str))
        """
        self.fichier_entree = nom_fichier
        questions = []
        
        try:
            with open(nom_fichier, 'r', encoding='utf-8') as fichier:
                contenu = fichier.read().strip().split('\n\n')
                
                for bloc in contenu:
                    lignes = bloc.strip().split('\n')
                    if len(lignes) >= 6:  # Au moins une question, 4 réponses et 1 solution
                        question = lignes[0]
                        reponses = lignes[1:5]
                        solution = int(lignes[5])
                        
                        questions.append({
                            'question': question,
                            'reponses': reponses,
                            'solution': solution
                        })
                
            self.questions = questions
            return True, f"{len(questions)} questions lues avec succès"
        except Exception as e:
            return False, f"Erreur lors de la lecture du fichier: {str(e)}"
    
    def generer_sujets(self, nombre_eleves, nombre_questions, graine):
        """
        Génère les sujets pour chaque élève.
        
        Args:
            nombre_eleves (int): Nombre de sujets à générer.
            nombre_questions (int): Nombre de questions par sujet.
            graine (int): Graine pour le générateur aléatoire.
            
        Returns:
            tuple: (succès (bool), sujets ou message d'erreur)
        """
        self.nombre_eleves = nombre_eleves
        self.nombre_questions = nombre_questions
        self.graine = graine
        
        random.seed(graine)  # Initialisation du générateur aléatoire
        
        if nombre_questions > len(self.questions):
            return False, (
                f"Pas assez de questions dans le fichier ({len(self.questions)}) "
                f"pour générer {nombre_questions} questions par sujet"
            )
        
        # Liste qui contiendra tous les sujets
        sujets = []
        
        # Pour chaque élève, on génère un sujet
        for num_sujet in range(nombre_eleves):
            # Sélection aléatoire des questions pour ce sujet
            questions_selectionnees = random.sample(self.questions, nombre_questions)
            
            # On prépare le sujet avec questions et réponses mélangées
            sujet = {
                'numero': num_sujet,
                'questions': []
            }
            
            for q in questions_selectionnees:
                # Copie des réponses pour les mélanger
                reponses = q['reponses'].copy()
                solution_originale = q['solution']
                
                # Mélange des réponses
                indices_melanges = list(range(4))
                random.shuffle(indices_melanges)
                
                reponses_melangees = [reponses[i] for i in indices_melanges]
                
                # Calculer la nouvelle position de la bonne réponse
                nouvelle_solution = indices_melanges.index(solution_originale - 1) + 1
                
                sujet['questions'].append({
                    'question': q['question'],
                    'reponses': reponses_melangees,
                    'solution': nouvelle_solution,
                    'question_originale_index': self.questions.index(q)
                })
            
            sujets.append(sujet)
        
        return True, sujets
    
    def generer_fichier_correction(self, sujets, nom_fichier="correction.txt"):
        """
        Génère le fichier de correction pour l'enseignant.
        
        Args:
            sujets (list): Liste des sujets générés.
            nom_fichier (str): Nom du fichier de correction à générer.
            
        Returns:
            tuple: (succès (bool), message (str))
        """
        try:
            with open(nom_fichier, 'w', encoding='utf-8') as f:
                for sujet in sujets:
                    f.write(f"Sujet {sujet['numero']}\n")
                    
                    # Écriture des réponses en groupes de 5
                    # Conversion 1->a, 2->b, etc.
                    reponses = [chr(96 + q['solution']) for q in sujet['questions']]
                    
                    for i in range(0, len(reponses), 5):
                        f.write(''.join(reponses[i:i+5]) + ' ')
                    
                    f.write('\n\n')
                
                # Approfondissement: traçabilité des questions
                f.write("\n--- Traçabilité des questions ---\n\n")
                
                # Pour chaque sujet, quelles questions initiales ont été intégrées
                for sujet in sujets:
                    indices_questions = [q['question_originale_index'] for q in sujet['questions']]
                    f.write(f"Sujet {sujet['numero']} : questions {','.join(map(str, indices_questions))}\n")
                
                # Pour chaque question initiale, dans quels sujets elle apparaît
                for i, _ in enumerate(self.questions):
                    sujets_avec_question = [
                        s['numero'] for s in sujets 
                        if any(q['question_originale_index'] == i for q in s['questions'])
                    ]
                    f.write(f"Question {i} : sujets {','.join(map(str, sujets_avec_question))}\n")
                
            return True, f"Fichier de correction généré avec succès: {nom_fichier}"
        except Exception as e:
            return False, f"Erreur lors de la génération du fichier de correction: {str(e)}"
    
    def generer_sujets_docx(self, sujets, dossier_sortie="sujets"):
        """
        Génère les sujets au format docx.
        
        Args:
            sujets (list): Liste des sujets générés.
            dossier_sortie (str): Dossier où seront enregistrés les fichiers.
            
        Returns:
            tuple: (succès (bool), message (str))
        """
        try:
            os.makedirs(dossier_sortie, exist_ok=True)
            
            for sujet in sujets:
                doc = Document()
                
                # Mise en page pour A4 paysage, 2 colonnes
                section = doc.sections[0]
                section.page_height = 210 * 36000  # A4 height in EMUs
                section.page_width = 297 * 36000   # A4 width in EMUs
                section.left_margin = 1 * 36000    # 1cm left margin
                section.right_margin = 1 * 36000   # 1cm right margin
                section.top_margin = 1 * 36000     # 1cm top margin
                section.bottom_margin = 1 * 36000  # 1cm bottom margin
                
                # En-tête du document
                doc.add_heading(f'QCM - Sujet {sujet["numero"]}', 0)
                doc.add_paragraph('Instructions : Pour chaque question, cochez la réponse qui vous semble correcte.')
                
                # Ajout des questions et réponses
                for i, q in enumerate(sujet['questions']):
                    doc.add_heading(f'Question {i+1} : {q["question"]}', level=2)
                    for j, rep in enumerate(q['reponses']):
                        doc.add_paragraph(f'{chr(97 + j)}) {rep}')
                
                # Ajout du cartouche pour les réponses
                doc.add_heading('Réponses', level=1)
                doc.add_paragraph('Nom prénom : ____________________')
                doc.add_paragraph(f'Sujet {sujet["numero"]} Réponses')
                
                # Tableau pour les réponses
                table = doc.add_table(rows=2, cols=10)
                
                # Première ligne: numéros de 1 à 10
                for i in range(10):
                    cell = table.cell(0, i)
                    cell.text = str(i+1)
                
                # Deuxième ligne: espaces pour les réponses
                for i in range(10):
                    cell = table.cell(1, i)
                    cell.text = '_'
                
                # Si plus de 10 questions, on ajoute une deuxième partie
                if len(sujet['questions']) > 10:
                    doc.add_paragraph('')  # Espace
                    
                    table2 = doc.add_table(rows=2, cols=10)
                    
                    # Première ligne: numéros de 11 à 20
                    for i in range(10):
                        cell = table2.cell(0, i)
                        cell.text = str(i+11)
                    
                    # Deuxième ligne: espaces pour les réponses
                    for i in range(10):
                        cell = table2.cell(1, i)
                        cell.text = '_'
                
                # Sauvegarde du document
                doc.save(f'{dossier_sortie}/sujet_{sujet["numero"]}.docx')
            
            return True, f"Les sujets ont été générés avec succès dans le dossier: {dossier_sortie}"
        except Exception as e:
            return False, f"Erreur lors de la génération des sujets: {str(e)}"

    def lancer_interface_graphique(self):
        """Lance l'interface graphique avec tkinter améliorée."""
        root = tk.Tk()
        root.title("Générateur de QCM")
        root.geometry("700x500")
        root.configure(bg="#f0f0f0")
        
        # Définition des styles et polices
        title_font = Font(family="Helvetica", size=16, weight="bold")
        section_font = Font(family="Helvetica", size=12, weight="bold")
        normal_font = Font(family="Helvetica", size=10)
        
        # Utilisation du style ttk pour un look plus moderne
        style = ttk.Style()
        style.configure("TFrame", background="#f0f0f0")
        style.configure("TLabel", background="#f0f0f0", font=normal_font)
        style.configure("TButton", font=normal_font)
        style.configure("Header.TLabel", font=section_font)
        style.configure("Title.TLabel", font=title_font)
        
        # Variables
        fichier_var = tk.StringVar()
        eleves_var = tk.IntVar(value=10)
        questions_var = tk.IntVar(value=5)
        graine_var = tk.IntVar(value=42)
        status_var = tk.StringVar(value="En attente d'un fichier QCM...")
        
        # Conteneur principal avec padding
        main_frame = ttk.Frame(root, padding="20 20 20 20", style="TFrame")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Titre de l'application
        title_label = ttk.Label(main_frame, text="Générateur de QCM", style="Title.TLabel")
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20), sticky="w")
        
        # Séparateur
        separator = ttk.Separator(main_frame, orient="horizontal")
        separator.grid(row=1, column=0, columnspan=3, sticky="ew", pady=(0, 20))
        
        # Section Fichier
        fichier_label = ttk.Label(main_frame, text="1. Sélection du fichier QCM", style="Header.TLabel")
        fichier_label.grid(row=2, column=0, columnspan=3, sticky="w", pady=(0, 10))
        
        help_text = ttk.Label(
            main_frame, 
            text="Sélectionnez un fichier texte (.txt) contenant les questions au format spécifié."
        )
        help_text.grid(row=3, column=0, columnspan=3, sticky="w", pady=(0, 10))
        
        fichier_frame = ttk.Frame(main_frame, style="TFrame")
        fichier_frame.grid(row=4, column=0, columnspan=3, sticky="ew", pady=(0, 20))
        
        fichier_entry = ttk.Entry(fichier_frame, textvariable=fichier_var, width=50)
        fichier_entry.pack(side=tk.LEFT, padx=(0, 10), fill=tk.X, expand=True)
        
        # Fonction pour sélectionner le fichier
        def selectionner_fichier():
            """Permet à l'utilisateur de sélectionner un fichier QCM."""
            fichier = filedialog.askopenfilename(
                title="Sélectionner un fichier QCM",
                filetypes=[("Fichiers texte", "*.txt")]
            )
            if fichier:
                fichier_var.set(fichier)
                
                # Lire le fichier pour afficher les infos
                success, message = self.lire_fichier_qcm(fichier)
                if success:
                    status_var.set(message)
                    messagebox.showinfo("Information", message)
                else:
                    status_var.set("Erreur: " + message)
                    messagebox.showerror("Erreur", message)
        
        fichier_button = ttk.Button(fichier_frame, text="Parcourir...", command=selectionner_fichier)
        fichier_button.pack(side=tk.LEFT)
        
        # Section Paramètres
        param_label = ttk.Label(main_frame, text="2. Paramètres de génération", style="Header.TLabel")
        param_label.grid(row=5, column=0, columnspan=3, sticky="w", pady=(0, 10))
        
        # Frame pour les paramètres avec 3 colonnes
        param_frame = ttk.Frame(main_frame, style="TFrame")
        param_frame.grid(row=6, column=0, columnspan=3, sticky="ew", pady=(0, 20))
        param_frame.columnconfigure(1, weight=1)
        param_frame.columnconfigure(3, weight=1)
        param_frame.columnconfigure(5, weight=1)
        
        # Nombre d'élèves
        ttk.Label(param_frame, text="Nombre d'élèves:").grid(
            row=0, column=0, sticky="w", padx=(0, 10), pady=10
        )
        eleves_spinbox = ttk.Spinbox(param_frame, from_=1, to=100, textvariable=eleves_var, width=10)
        eleves_spinbox.grid(row=0, column=1, sticky="w", padx=10, pady=10)
        
        # Nombre de questions
        ttk.Label(param_frame, text="Nombre de questions:").grid(
            row=0, column=2, sticky="w", padx=10, pady=10
        )
        questions_combo = ttk.Combobox(
            param_frame, 
            textvariable=questions_var, 
            values=[5, 10, 15, 20], 
            width=10, 
            state="readonly"
        )
        questions_combo.grid(row=0, column=3, sticky="w", padx=10, pady=10)
        
        # Graine aléatoire
        ttk.Label(param_frame, text="Graine aléatoire:").grid(
            row=0, column=4, sticky="w", padx=10, pady=10
        )
        graine_spinbox = ttk.Spinbox(param_frame, from_=0, to=9999, textvariable=graine_var, width=10)
        graine_spinbox.grid(row=0, column=5, sticky="w", padx=(10, 0), pady=10)
        
        # Section Génération
        gen_label = ttk.Label(main_frame, text="3. Génération des QCM", style="Header.TLabel")
        gen_label.grid(row=7, column=0, columnspan=3, sticky="w", pady=(0, 10))
        
        # Fonction pour générer les QCM
        def generer_qcm():
            """Génère les QCM selon les paramètres spécifiés."""
            if not self.questions:
                messagebox.showerror("Erreur", "Veuillez d'abord charger un fichier de questions.")
                return
            
            n_eleves = eleves_var.get()
            n_questions = questions_var.get()
            graine = graine_var.get()
            
            if n_questions not in [5, 10, 15, 20]:
                messagebox.showerror(
                    "Erreur", 
                    "Le nombre de questions doit être un multiple de 5 (5, 10, 15 ou 20)."
                )
                return
            
            status_var.set("Génération des sujets en cours...")
            root.update_idletasks()
            
            # Génération des sujets
            success, sujets_ou_message = self.generer_sujets(n_eleves, n_questions, graine)
            
            if not success:
                status_var.set("Erreur: " + sujets_ou_message)
                messagebox.showerror("Erreur", sujets_ou_message)
                return
            
            sujets = sujets_ou_message
            
            # Demande où sauvegarder les fichiers
            dossier_sortie = filedialog.askdirectory(title="Sélectionnez le dossier de sortie")
            if not dossier_sortie:
                status_var.set("Génération annulée par l'utilisateur.")
                return
            
            status_var.set("Génération des fichiers de correction...")
            root.update_idletasks()
            
            # Génération du fichier de correction
            success_correction, message_correction = self.generer_fichier_correction(
                sujets, f"{dossier_sortie}/correction.txt"
            )
            
            status_var.set("Génération des sujets DOCX...")
            root.update_idletasks()
            
            # Génération des sujets au format docx
            success_docx, message_docx = self.generer_sujets_docx(sujets, dossier_sortie)
            
            if success_correction and success_docx:
                status_var.set("Génération terminée avec succès!")
                messagebox.showinfo(
                    "Succès", 
                    f"Génération des QCM terminée.\n\n{message_correction}\n\n{message_docx}"
                )
            else:
                error_msg = "Erreurs lors de la génération des QCM:\n"
                if not success_correction:
                    error_msg += message_correction + "\n"
                if not success_docx:
                    error_msg += message_docx
                status_var.set("Erreur lors de la génération!")
                messagebox.showerror("Erreur", error_msg)
        
        # Frame pour le bouton et la barre de statut
        action_frame = ttk.Frame(main_frame, style="TFrame")
        action_frame.grid(row=8, column=0, columnspan=3, sticky="ew", pady=(0, 20))
        
        generate_button = ttk.Button(
            action_frame, 
            text="Générer les QCM", 
            command=generer_qcm,
            padding=(20, 10)
        )
        generate_button.pack(pady=10)
        
        # Séparateur
        separator2 = ttk.Separator(main_frame, orient="horizontal")
        separator2.grid(row=9, column=0, columnspan=3, sticky="ew", pady=10)
        
        # Barre de statut
        status_frame = ttk.Frame(main_frame, style="TFrame")
        status_frame.grid(row=10, column=0, columnspan=3, sticky="ew")
        
        status_label = ttk.Label(status_frame, text="Statut:")
        status_label.pack(side=tk.LEFT, padx=(0, 10))
        
        status_text = ttk.Label(status_frame, textvariable=status_var, foreground="#555555")
        status_text.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Aide contextuelle
        help_frame = ttk.Frame(main_frame, padding="10", relief="groove", borderwidth=1)
        help_frame.grid(row=11, column=0, columnspan=3, sticky="ew", pady=(20, 0))
        
        help_title = ttk.Label(help_frame, text="Aide", font=section_font)
        help_title.pack(anchor="w")
        
        help_content = ttk.Label(
            help_frame, 
            text="• Le fichier QCM doit être au format texte (.txt)\n"
                "• Chaque question doit suivre le format: question, 4 réponses, numéro de la bonne réponse\n"
                "• Une ligne vide doit séparer chaque question\n"
                "• Le nombre de questions par sujet doit être un multiple de 5 (5, 10, 15, 20)",
            justify="left"
        )
        help_content.pack(anchor="w", pady=(5, 0))
        
        root.mainloop()


if __name__ == "__main__":
    generateur = GenerateurQCM()
    generateur.lancer_interface_graphique()